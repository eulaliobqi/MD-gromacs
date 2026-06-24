#!/usr/bin/env python3
"""
embed_figuras_docx.py — Insere as figuras compostas em artigo_md.docx.

Lê o artigo_md.docx existente, localiza os marcadores "[[FIG_N]]" no texto,
e substitui cada marcador pela imagem Figure_N.png + legenda.

Uso (na raiz do repositório):
  python3 bin/embed_figuras_docx.py [--figures figures/] [--input artigo_md.docx] [--output artigo_md_figs.docx]

Pré-requisito: rodar gerar_figuras_compostas.py antes para gerar as PNGs.
"""
import argparse
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERRO: pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).parent.parent

# Legendas completas de cada figura
FIGURE_CAPTIONS = {
    1: (
        "Figura 1 — Parâmetros dinâmicos dos complexos peptídeo GORE4 × tripsinas "
        "de Spodoptera (100 ns, 300 K, pH 8,2). "
        "(A) QCL936-GORE4: RMSD backbone = 0,124 ± 0,017 nm, estável após 20 ns; "
        "RMSD ligante = 0,229 ± 0,042 nm; contatos = 340 ± 41 átomos; H-bonds = 2,75 ± 1,01. "
        "(B) ACR157-GORE4: RMSD backbone = 0,165 nm; RMSD ligante = 0,393 nm; "
        "contatos = 256 ± 38 com reorganização aos 30 ns. "
        "(C) XP273-GORE4: perfil bimodal do raio de giro (1,727 ± 0,023 nm) indica "
        "dois estados conformacionais; RMSD ligante = 0,317 nm; contatos = 260 ± 60. "
        "(D) XP352-GORE4 (controle de instabilidade): RMSD backbone = 0,886 ± 0,440 nm "
        "sem platô; contatos = 63 ± 143 (colapso diagnóstico). "
        "Médias móveis calculadas com janela de 5 ns."
    ),
    2: (
        "Figura 2 — Distâncias mínimas entre o peptídeo GORE4 e os resíduos do sítio "
        "catalítico de cada receptor ao longo de 100 ns. Linha tracejada: limiar de "
        "contato direto (0,5 nm). triad_1 = His/Tyr/Arg; triad_2 = Asp catalítico; "
        "triad_3 = Ser nucleofílica; triad_4 = resíduo do bolsão S1. "
        "(A) QCL936-GORE4: triad_1 (His92, 0,242 nm) e triad_4 (Asp241, 0,227 nm) "
        "abaixo do limiar — modo His+S1. "
        "(B) ACR157-GORE4: triad_1 (His69) borderline; triad_4 (Ile205) parcial. "
        "(C) XP273-GORE4: apenas triad_1 (Tyr83) abaixo de 0,5 nm — modo periférico. "
        "(D) XP352-GORE4: todos os resíduos > 1,0 nm — dissociação confirmada."
    ),
    3: (
        "Figura 3 — Parâmetros dinâmicos dos complexos SKTI × tripsinas de Spodoptera "
        "(100 ns, 300 K, pH 8,2). "
        "(A) ACR157-SKTI: RMSD ligante = 0,206 ± 0,017 nm (menor de toda a série); "
        "contatos = 1019 ± 118; H-bonds = 13,524 ± 2,911 — interface máxima. "
        "(B) QCL936-SKTI: RMSD backbone = 0,370 ± 0,054 nm (maior mobilidade do receptor "
        "com SKTI); contatos = 720 ± 77. "
        "(C) XP273-SKTI: RMSD ligante = 0,224 ± 0,017 nm; contatos = 596 ± 88; "
        "H-bonds = 8,817 ± 3,496. SASA ligante ≈ 101–103 nm² nos três sistemas "
        "confirma enterramento estável do SKTI na interface."
    ),
    4: (
        "Figura 4 — Distâncias mínimas entre SKTI e os resíduos do sítio catalítico "
        "(mesma nomenclatura da Figura 2). "
        "(A) ACR157-SKTI: mecanismo Kunitz canônico completo — todos os quatro resíduos "
        "(His69, Asp114, Ser211, Ile205/S1) abaixo de 0,35 nm durante toda a trajetória. "
        "(B) QCL936-SKTI: modo His+Ser+S1 — His92 (0,179 nm), Ser247 (0,176 nm) e "
        "Asp241 (0,178 nm) em contato; Asp142 distante (~0,73 nm). "
        "(C) XP273-SKTI: modo Tyr+Asp+S1 — Tyr83 (0,098 nm), Asp132 (0,119 nm) e "
        "Ile229 (0,118 nm) engajados; Ser234 = 0,718 nm."
    ),
    5: (
        "Figura 5 — Evento de dissociação da benzamidina (BEN) em quatro isoformas de "
        "tripsina de Spodoptera (200 ns). Painéis organizados em ordem crescente de "
        "tempo de residência. (A) XP273-BEN (S1=Ile229 neutro): dissociação a ~80 ns; "
        "ausência de âncora eletrostática. "
        "(B) ACR157-BEN (S1=Ile205 neutro): dissociação a ~95 ns; "
        "mesma ausência de ponte salina. "
        "(C) XP352-BEN (S1=Asp262, −1): ancoragem borderline 0–125 ns; dissociação "
        "a ~125 ns. "
        "(D) QCL936-BEN (S1=Asp241, −1): fase ligada robusta 0–150 ns com His92, "
        "Ser247 e Asp241 em contato; dissociação a ~150 ns. "
        "Linhas tracejadas indicam o instante de dissociação. BEN: benzamidina "
        "(CAS 618-39-3; GAFF2/ACPYPE, carga +1)."
    ),
    6: (
        "Figura 6 — Mapas de frequência de contato resíduo×resíduo (distância < 0,4 nm) "
        "para os seis complexos estáveis (stride 10, MDAnalysis 2.10). Eixo x: resíduos "
        "do ligante; eixo y: resíduos do receptor. Intensidade de cor = fração de frames "
        "com ao menos um par atômico a < 0,4 nm. Agrupamento hierárquico Ward (Seaborn "
        "clustermap). (A–C) Série GORE4: QCL936, ACR157 e XP273. "
        "(D–F) Série SKTI: ACR157, QCL936 e XP273. "
        "Notar maior cobertura de interface para SKTI (maior extensão do mapa) "
        "versus GORE4 (padrão linear de 5 resíduos)."
    ),
    7: (
        "Figura 7 — Fingerprints de interação molecular ProLIF: persistência temporal "
        "de cada par resíduo(ligante)–resíduo(receptor)–tipo de interação (%). "
        "Análise realizada com ProLIF 2.x (Bouysset & Fiorucci, 2021), stride 10 (1 ns/frame). "
        "Tipos: HBDonor/HBAcceptor (ligações de hidrogênio), Hydrophobic, Catiônica, Aniônica, "
        "VdWContact. (A–C) Série GORE4: padrão decrescente de persistência "
        "QCL936 (100% Catiônica LYS303-ALA242) > ACR157 (52% VdW ALA256-GLY228) > "
        "XP273 (28% Catiônica LEU280-Tyr89). "
        "(D–F) Série SKTI: ARG317-GLN206 100% em ACR157 (Kunitz canônico); "
        "ARG361/ARG363 bipartite 100% em QCL936; ARG344-HID86 58% em XP273 "
        "(His catalítica — novo achado)."
    ),
}

# Marcadores no texto que serão substituídos pelas figuras
# Formato esperado no docx: parágrafo contendo exatamente [[FIG_N]]
MARKERS = {f"[[FIG_{n}]]": n for n in range(1, 8)}


def insert_figure(doc, para_idx, fig_path, caption_text, fig_width_inches=6.3):
    """Substitui o parágrafo marcador pela imagem + legenda."""
    para = doc.paragraphs[para_idx]
    # Limpa o parágrafo marcador
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insere imagem
    run = para.add_run()
    if fig_path and fig_path.exists():
        run.add_picture(str(fig_path), width=Inches(fig_width_inches))
    else:
        run.text = f"[Figura não encontrada: {fig_path}]"
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # Legenda
    cap_para = para.insert_paragraph_before("")
    # Mover a imagem para antes da legenda: trocar ordem
    # Na prática em python-docx é mais fácil: adicionar legenda depois
    cap_para2 = OxmlElement("w:p")
    # Adicionar legenda logo após o parágrafo de imagem
    new_para = doc.add_paragraph(caption_text)
    new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para.style.font.size = Pt(9)
    new_para.style.font.italic = True
    # Mover o parágrafo para após o da figura
    para._p.addnext(new_para._p)


def process_doc(input_path, output_path, figures_dir):
    doc = Document(str(input_path))
    figures_dir = Path(figures_dir)
    replaced = 0

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text in MARKERS:
            n = MARKERS[text]
            fig_path = figures_dir / f"Figure_{n}.png"
            caption = FIGURE_CAPTIONS.get(n, f"Figura {n}.")
            print(f"  Inserindo Figura {n} no parágrafo {i} ('{text}')")
            insert_figure(doc, i, fig_path, caption)
            replaced += 1

    print(f"  {replaced} marcadores substituídos.")
    doc.save(str(output_path))
    print(f"  Salvo: {output_path}")


def main():
    parser = argparse.ArgumentParser(description=__doc__,
                                     formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--figures", default="figures",
                        help="Diretório com Figure_N.png (default: figures/)")
    parser.add_argument("--input", default="artigo_md.docx",
                        help="DOCX de entrada (default: artigo_md.docx)")
    parser.add_argument("--output", default="artigo_md_figs.docx",
                        help="DOCX de saída com figuras (default: artigo_md_figs.docx)")
    args = parser.parse_args()

    input_path  = ROOT / args.input
    output_path = ROOT / args.output
    figures_dir = ROOT / args.figures

    if not input_path.exists():
        print(f"ERRO: {input_path} não encontrado. Rodar gerar_artigo_docx.py primeiro.")
        sys.exit(1)

    print(f"Input  : {input_path}")
    print(f"Figuras: {figures_dir}")
    print(f"Output : {output_path}")
    print()

    process_doc(input_path, output_path, figures_dir)
    print()
    print("[OK] embed_figuras_docx.py concluído")


if __name__ == "__main__":
    main()
