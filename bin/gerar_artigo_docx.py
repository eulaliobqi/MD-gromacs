"""
Gera artigo_md.docx a partir de artigo_md.md
Usa python-docx com formatação científica (Times New Roman, 12pt, espaçamento 1.5)
"""
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Instalar: pip install python-docx")
    sys.exit(1)


SOURCE = Path(__file__).parent.parent / "artigo_md.md"
OUTPUT = Path(__file__).parent.parent / "artigo_md.docx"


def set_cell_bg(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_paragraph(doc, text: str, style='Normal', bold=False, italic=False, align=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def apply_inline(para, text: str):
    """Processa **bold** e *italic* inline."""
    # Padrão: **texto** ou *texto*
    pattern = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`')
    last = 0
    for m in pattern.finditer(text):
        # texto antes do match
        before = text[last:m.start()]
        if before:
            para.add_run(before)
        if m.group(1) is not None:
            r = para.add_run(m.group(1))
            r.bold = True
        elif m.group(2) is not None:
            r = para.add_run(m.group(2))
            r.italic = True
        elif m.group(3) is not None:
            r = para.add_run(m.group(3))
            r.font.name = 'Courier New'
            r.font.size = Pt(10)
        last = m.end()
    remainder = text[last:]
    if remainder:
        para.add_run(remainder)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        if re.match(r'\|[-: |]+\|', line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    return rows


def create_word(md_path: Path, out_path: Path):
    doc = Document()

    # Margens (2,5 cm ABNT-like)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # Estilo Normal base
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    para_fmt = style.paragraph_format
    para_fmt.space_after = Pt(6)
    para_fmt.line_spacing = Pt(18)  # ~1.5

    # Estilos de cabeçalho
    for lvl, (name, size, bold) in enumerate([
        ('Heading 1', 16, True),
        ('Heading 2', 14, True),
        ('Heading 3', 13, True),
        ('Heading 4', 12, True),
    ], start=1):
        h = doc.styles[name]
        h.font.name = 'Times New Roman'
        h.font.size = Pt(size)
        h.font.bold = bold
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

    text = md_path.read_text(encoding='utf-8')
    lines = text.splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]

        # Título H1
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # H2
        if line.startswith('## ') and not line.startswith('### '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # H3
        if line.startswith('### ') and not line.startswith('#### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # H4
        if line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue

        # Separador horizontal
        if line.strip() in ('---', '***', '___'):
            doc.add_paragraph('', style='Normal')
            i += 1
            continue

        # Comentário HTML (PLACEHOLDER)
        if line.strip().startswith('<!--'):
            while i < len(lines) and '-->' not in lines[i]:
                i += 1
            i += 1
            continue

        # Linha de tabela: acumular linhas da tabela
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            if not rows:
                continue
            ncols = len(rows[0])
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = 'Table Grid'
            for r_idx, row in enumerate(rows):
                for c_idx, cell_text in enumerate(row):
                    cell = table.cell(r_idx, c_idx)
                    # limpar bold markdown
                    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                    clean = re.sub(r'\*(.+?)\*', r'\1', clean)
                    # remover emojis problemáticos
                    clean = re.sub(r'[✅❌⚠️⏳]', '', clean).strip()
                    cell.paragraphs[0].clear()
                    run = cell.paragraphs[0].add_run(clean)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    if r_idx == 0:
                        run.bold = True
                        set_cell_bg(cell, 'D9E1F2')
            doc.add_paragraph('', style='Normal')
            continue

        # Linha vazia
        if not line.strip():
            i += 1
            continue

        # Lista com bullet (- item)
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            content = line.strip()[2:]
            apply_inline(p, content)
            i += 1
            continue

        # Lista numerada
        m_num = re.match(r'^(\d+)\. (.+)', line.strip())
        if m_num:
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            apply_inline(p, m_num.group(2))
            i += 1
            continue

        # Bloco de código
        if line.strip().startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph('\n'.join(code_lines), style='Normal')
                p.runs[0].font.name = 'Courier New'
                p.runs[0].font.size = Pt(9)
                p.paragraph_format.left_indent = Cm(1)
            i += 1
            continue

        # Nota de rodapé-like (ª ᵇ ᶜ)
        if re.match(r'^[ªᵃᵇᶜᵈ]', line.strip()):
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.space_before = Pt(0)
            apply_inline(p, line.strip())
            p.runs[0].font.size = Pt(9)
            i += 1
            continue

        # Parágrafo normal
        p = doc.add_paragraph(style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        apply_inline(p, line.strip())
        i += 1

    doc.save(out_path)
    print(f"[OK] Arquivo salvo: {out_path}")


if __name__ == '__main__':
    create_word(SOURCE, OUTPUT)
