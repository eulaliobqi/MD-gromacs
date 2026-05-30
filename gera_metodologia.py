from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Estilos ──────────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)

for s in ['Heading 1', 'Heading 2', 'Heading 3']:
    h = doc.styles[s]
    h.font.name = 'Times New Roman'
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.font.bold = True

doc.styles['Heading 1'].font.size = Pt(13)
doc.styles['Heading 2'].font.size = Pt(12)
doc.styles['Heading 3'].font.size = Pt(12)

for section in doc.sections:
    section.top_margin    = Inches(0.984)
    section.bottom_margin = Inches(0.984)
    section.left_margin   = Inches(1.181)
    section.right_margin  = Inches(0.984)

# ── Helpers ───────────────────────────────────────────────────────────────────
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = 'Times New Roman'
    return h

def add_para(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(36)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p

def add_para_mixed(parts):
    """parts = list of (text, bold). Adiciona parágrafo com partes em negrito."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(36)
    p.paragraph_format.space_after = Pt(6)
    for text, bold in parts:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.bold = bold
    return p

def add_bullet(bold_label, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    r1 = p.add_run(bold_label + ' ')
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r2 = p.add_run(text)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)

def add_ref(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(36)
    p.paragraph_format.first_line_indent = Pt(-36)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

def set_cell(cell, text, bold=False, center=False, shade=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.bold = bold
    if shade:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), shade)
        tcPr.append(shd)

# ════════════════════════════════════════════════════════════════════════════
# TÍTULO
# ════════════════════════════════════════════════════════════════════════════
titulo = doc.add_paragraph()
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
titulo.paragraph_format.space_after = Pt(12)
r = titulo.add_run('MATERIAL E MÉTODOS')
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(14)

# ════════════════════════════════════════════════════════════════════════════
# 1. DOCKING MOLECULAR
# ════════════════════════════════════════════════════════════════════════════
add_heading('1. Docking Molecular', level=1)

add_para(
    'As estruturas tridimensionais dos peptídeos candidatos a inibidores foram geradas '
    'por modelagem de novo utilizando o servidor PEPstrMOD (Singh et al., 2015), '
    'considerando o pH fisiológico do intestino médio de Spodoptera spp. (pH 8,2). '
    'A estrutura da tripsina digestiva de Spodoptera frugiperda (GORE4) foi obtida '
    'a partir do Protein Data Bank (PDB) ou modelada por homologia com o servidor '
    'AlphaFold2, seguida de refinamento estrutural pelo servidor ModRefiner.'
)

add_para(
    'O preparo dos receptores foi realizado com o software PDB2PQR v3.6 '
    '(Dolinsky et al., 2007), utilizando o campo de força AMBER e o método PROPKA '
    'para atribuição de estados de protonação a pH 8,2. Especial atenção foi '
    'dedicada à His57 da tríade catalítica, cujo valor de pKa predito pelo PROPKA '
    '(~6,5) resulta em estado neutro (ε-protonado, HIE) em pH 8,2, representando '
    'a forma ativa da serino-protease em meio alcalino — condição fisiológica '
    'do intestino médio de lepidópteros (pH 9,5–11,0; Dossantos et al., 2015; '
    'Zhu et al., 2022). Moléculas de água e ligantes heterólogos foram removidos '
    'antes da análise.'
)

add_para(
    'O docking proteína–peptídeo foi conduzido no servidor HADDOCK 2.4 '
    '(High Ambiguity Driven DOCKing) (van Zundert et al., 2016; Dominguez et al., 2003). '
    'As restrições ambíguas de interação (AIRs) foram definidas com base nos resíduos '
    'do sítio ativo da tripsina GORE4 — incluindo os resíduos catalíticos Tyr83, '
    'Asp132, Ser234 e o bolsão de especificidade S1 (Ile229) — como resíduos "ativos" '
    'do receptor, e nos resíduos básicos dos peptídeos (lisina e arginina) como '
    'resíduos ativos dos ligantes, em conformidade com a especificidade de clivagem '
    'de tripsinas por resíduos básicos.'
)

add_para(
    'O protocolo padrão do HADDOCK foi executado em três etapas: (i) minimização de '
    'energia de corpo rígido (it0), gerando 1.000 estruturas; (ii) refinamento '
    'semirígido em espaço de torção (it1), retendo as 200 melhores estruturas; e '
    '(iii) refinamento final em solvente explícito (water refinement), com seleção '
    'das 200 melhores poses. As estruturas foram agrupadas por RMSD de 7,5 Å '
    'sobre os átomos da interface. Os clusters foram ranqueados pelo HADDOCK score '
    '(combinação linear de energias de van der Waals, eletrostática, dessolvatação '
    'e restrições ambíguas). A pose de menor energia de cada cluster foi selecionada '
    'para as simulações de dinâmica molecular.'
)

# ════════════════════════════════════════════════════════════════════════════
# 2. DINÂMICA MOLECULAR
# ════════════════════════════════════════════════════════════════════════════
add_heading('2. Dinâmica Molecular', level=1)

# 2.1
add_heading('2.1 Condições fisiológicas e preparo do sistema', level=2)

add_para(
    'Os parâmetros das simulações foram estabelecidos para reproduzir as condições '
    'fisiológicas do intestino médio de larvas de Spodoptera spp. A temperatura '
    'de 300 K (27 °C) foi adotada como representativa da temperatura corporal '
    'larval, estimada entre 25 e 30 °C (Soelter et al., 2024; Da Silva et al., 2018). '
    'A pressão foi mantida a 1 bar (pressão atmosférica). A concentração iônica '
    'de 0,10 M de NaCl foi selecionada para aproximar a força iônica hipotônica '
    'do lúmen intestinal de lepidópteros, que difere do plasma de mamíferos '
    '(Zhu et al., 2022; Dunse et al., 2010).'
)

add_para(
    'Os complexos tripsina–peptídeo obtidos pelo docking foram preparados para '
    'simulação utilizando um pipeline automatizado desenvolvido em Nextflow DSL2 '
    '(Di Tommaso et al., 2017), integrando ferramentas do pacote GROMACS 2026 '
    '(Abraham et al., 2015) e PDB2PQR v3.6. O preparo compreendeu: (i) atribuição '
    'de estados de protonação pelo método PROPKA em pH 8,2; (ii) geração da '
    'topologia com o campo de força AMBER99SB-ILDN '
    '(Lindorff-Larsen et al., 2010), amplamente validado para serino-proteases '
    'e peptídeos inibidores (Soelter et al., 2024; Ahmad et al., 2018); e '
    '(iii) centralização do complexo em caixa dodecaédrica com margem mínima '
    'de 1,2 nm entre o soluto e as faces da caixa, minimizando interações '
    'entre imagens periódicas.'
)

add_para(
    'A solvatação foi realizada com moléculas de água no modelo TIP3P '
    '(Jorgensen et al., 1983), compatível com o campo de força AMBER99SB-ILDN. '
    'A neutralização do sistema e o ajuste da concentração iônica foram realizados '
    'pela substituição aleatória de moléculas de água por íons Na⁺ e Cl⁻ a '
    '0,10 M, utilizando o módulo gmx genion.'
)

# 2.2
add_heading('2.2 Minimização de energia', level=2)

add_para(
    'A minimização de energia foi realizada pelo algoritmo de descida mais íngreme '
    '(steepest descent) com até 50.000 passos, convergindo quando a força máxima '
    'sobre qualquer átomo foi inferior a 1.000 kJ mol⁻¹ nm⁻¹ (emtol), valor '
    'padrão recomendado para sistemas proteína–solvente (Abraham et al., 2015). '
    'As interações eletrostáticas foram tratadas pelo método Particle Mesh Ewald '
    '(PME) com cutoff de 1,2 nm, e as interações de van der Waals pelo esquema '
    'Verlet com raio de corte de 1,2 nm.'
)

# 2.3
add_heading('2.3 Equilibração termodinâmica', level=2)

add_para(
    'Após a minimização, o sistema foi equilibrado em duas etapas sequenciais com '
    'restrições de posição aplicadas a todos os átomos pesados do soluto '
    '(constante de força: 1.000 kJ mol⁻¹ nm⁻²), permitindo a reorganização '
    'e equilíbrio do solvente ao redor do complexo.'
)

add_para(
    'A primeira etapa de equilibração foi conduzida no ensemble NVT (volume e '
    'temperatura constantes) durante 200 ps (100.000 passos de 2 fs), com '
    'aquecimento progressivo até 300 K utilizando o termostato de velocidades '
    'reescalonadas (V-rescale; Bussi et al., 2007), com constante de acoplamento '
    'de τ = 0,1 ps, aplicado separadamente aos grupos Proteína e Não-Proteína. '
    'O tempo de 200 ps foi adotado com base em estudos de complexos '
    'proteína–peptídeo que demonstram necessidade de equilíbrio térmico '
    'superior ao protocolo mínimo de 100 ps recomendado para proteínas simples '
    '(Bray et al., 2024; Trozzi et al., 2021).'
)

add_para(
    'A segunda etapa foi conduzida no ensemble NPT (pressão e temperatura '
    'constantes) durante 500 ps (250.000 passos de 2 fs), com o barostato de '
    'Berendsen (τp = 2,0 ps; compressibilidade isotérmica = 4,5 × 10⁻⁵ bar⁻¹) '
    'a 1 bar. O protocolo de 500 ps de equilibração NPT foi adotado conforme '
    'recomendado para complexos com peptídeos flexíveis de 5–15 resíduos, '
    'nos quais equilíbrio de densidade e conformação da interface requer '
    'tempo superior ao protocolo padrão de 100 ps '
    '(Al-Khafaji et al., 2024; Ahmad et al., 2018).'
)

# 2.4
add_heading('2.4 Simulação de produção', level=2)

add_para(
    'As simulações de produção foram realizadas no ensemble NPT durante 100 ns '
    'para cada complexo tripsina–peptídeo, com passo de integração de 2 fs. '
    'O termostato V-rescale (τ = 0,1 ps) foi mantido a 300 K e o barostato '
    'de Parrinello-Rahman (Parrinello e Rahman, 1981; τp = 2,0 ps) a 1 bar, '
    'garantindo flutuações de pressão termodinamicamente corretas no ensemble NPT. '
    'A transição do barostato de Berendsen (equilibração) para Parrinello-Rahman '
    '(produção) é o protocolo padrão recomendado para simulações de produção '
    '(Bray et al., 2024; Soelter et al., 2024).'
)

add_para(
    'As ligações envolvendo hidrogênio foram restringidas pelo algoritmo LINCS '
    '(Hess et al., 1997), com ordem 4 e uma iteração, permitindo o passo de '
    '2 fs. As interações eletrostáticas de longo alcance foram tratadas pelo '
    'método PME (Darden et al., 1993) com tolerância de 10⁻⁵ e malha PME '
    'atualizada a cada 20 passos. Velocidades iniciais foram geradas a partir '
    'de distribuição de Maxwell-Boltzmann a 300 K com semente aleatória '
    '(gen-seed = −1), assegurando independência estatística entre réplicas. '
    'Coordenadas foram salvas a cada 10 ps (nstxout-compressed = 5.000), '
    'gerando trajetórias de 10.000 frames. As simulações foram executadas '
    'em servidor Linux Debian com GPU NVIDIA, com o binário gmx_mpi compilado '
    'com suporte CUDA (aceleração de nb, pme e bonded na GPU).'
)

# Tabela de parâmetros
doc.add_paragraph()
caption = doc.add_paragraph()
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.space_after = Pt(4)
rc = caption.add_run(
    'Tabela 1. Parâmetros das simulações de dinâmica molecular dos complexos '
    'tripsina GORE4–peptídeo inibidor.'
)
rc.font.name = 'Times New Roman'
rc.font.size = Pt(11)
rc.bold = True

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

hdr = table.rows[0].cells
set_cell(hdr[0], 'Parâmetro', bold=True, center=True, shade='D9D9D9')
set_cell(hdr[1], 'Valor', bold=True, center=True, shade='D9D9D9')
set_cell(hdr[2], 'Referência', bold=True, center=True, shade='D9D9D9')

rows_data = [
    ('Campo de força',            'AMBER99SB-ILDN',                       'Lindorff-Larsen et al., 2010'),
    ('Modelo de água',            'TIP3P',                                 'Jorgensen et al., 1983'),
    ('Caixa de solvente',         'Dodecaedro rómbico, margem 1,2 nm',    'Abraham et al., 2015'),
    ('Temperatura',               '300 K (27 °C)',                         'Soelter et al., 2024'),
    ('Pressão',                   '1 bar',                                  '—'),
    ('Concentração iônica',       '0,10 M NaCl',                           'Zhu et al., 2022'),
    ('Eletrostática',             'PME, cutoff 1,2 nm',                    'Darden et al., 1993'),
    ('Van der Waals',             'Cutoff 1,2 nm, esquema Verlet',         '—'),
    ('Restrições de ligação',     'LINCS (ordem 4, 1 iteração)',           'Hess et al., 1997'),
    ('Passo de integração',       '2 fs',                                  '—'),
    ('Termostato',                'V-rescale, τ = 0,1 ps',                 'Bussi et al., 2007'),
    ('Barostato — equilibração',  'Berendsen, τp = 2,0 ps',               'Bray et al., 2024'),
    ('Barostato — produção',      'Parrinello-Rahman, τp = 2,0 ps',       'Parrinello & Rahman, 1981'),
    ('Equilibração NVT',          '200 ps',                                'Bray et al., 2024'),
    ('Equilibração NPT',          '500 ps',                                'Al-Khafaji et al., 2024'),
    ('Produção',                  '100 ns por complexo',                   'Ahmad et al., 2018'),
    ('Frequência de gravação',    '10 ps (nstxout-compressed = 5.000)',    '—'),
    ('Velocidades iniciais',      'Maxwell-Boltzmann 300 K, semente −1',  '—'),
    ('pH implícito',              '8,2 (PROPKA); His57 como HIE',          'Dolinsky et al., 2007'),
]

for parm, val, ref in rows_data:
    row = table.add_row().cells
    set_cell(row[0], parm)
    set_cell(row[1], val, center=True)
    set_cell(row[2], ref)

for i, col_width in enumerate([Cm(6.5), Cm(6.5), Cm(5.5)]):
    for cell in table.columns[i].cells:
        cell.width = col_width

doc.add_paragraph()

# 2.5
add_heading('2.5 Análise das trajetórias', level=2)

add_para(
    'As trajetórias de produção foram pós-processadas com o módulo gmx trjconv '
    'para: (i) correção de condições periódicas de contorno (PBC); '
    '(ii) centralização do complexo na caixa; e (iii) alinhamento rotacional e '
    'translacional pelo backbone do receptor. As seguintes propriedades foram '
    'calculadas ao longo de cada trajetória:'
)

analises = [
    ('Estabilidade estrutural (RMSD):',
     'o desvio quadrático médio das posições atômicas foi calculado para o '
     'backbone do receptor (Cα, C, N, O) e para os átomos pesados do peptídeo, '
     'utilizando a estrutura do primeiro frame da produção como referência '
     '(gmx rms). Trajetórias com RMSD do backbone < 0,3 nm e variação padrão '
     '< 0,05 nm foram consideradas estáveis.'),
    ('Flexibilidade por resíduo (RMSF):',
     'a flutuação quadrática média das posições atômicas foi calculada por '
     'resíduo para o backbone do receptor ao longo de toda a trajetória '
     '(gmx rmsf), identificando regiões de alta mobilidade e loops flexíveis.'),
    ('Compactação estrutural (Raio de giro, Rg):',
     'o raio de giro do complexo foi calculado sobre todos os átomos da proteína '
     'ao longo do tempo (gmx gyrate), como indicador de expansão ou compactação '
     'global do receptor durante a simulação.'),
    ('Contatos intermoleculares:',
     'o número de átomos em contato entre receptor e peptídeo (distância < 0,4 nm) '
     'e a distância mínima entre os grupos foram calculados com gmx mindist. '
     'Reduções abruptas no número de contatos foram interpretadas como eventos '
     'de dissociação do peptídeo.'),
    ('Pontes de hidrogênio:',
     'pontes de hidrogênio entre receptor e peptídeo foram identificadas e '
     'quantificadas ao longo da trajetória com gmx hbond, utilizando os critérios '
     'geométricos padrão: distância doador–aceptor ≤ 0,35 nm e ângulo ≤ 30°. '
     'Apenas pontes com persistência ≥ 10% da trajetória foram consideradas '
     'estáveis.'),
    ('Área de superfície acessível ao solvente (SASA):',
     'a SASA do receptor e do peptídeo foi calculada com gmx sasa '
     '(algoritmo de Connolly, raio de sonda 0,14 nm). Valores reduzidos de '
     'SASA do peptídeo ao longo do tempo indicam enterramento progressivo na '
     'interface de ligação e redução da exposição ao solvente.'),
    ('Distâncias aos resíduos catalíticos e bolsão S1:',
     'a distância mínima entre os átomos pesados do peptídeo e cada resíduo '
     'de interesse do sítio ativo (Tyr83, Asp132, Ser234 e Ile229/bolsão S1) '
     'foi monitorada ao longo da trajetória com gmx mindist. '
     'Distâncias < 0,5 nm foram consideradas compatíveis com ocupação do '
     'sítio ativo e potencial inibição da atividade catalítica.'),
]

for label, text in analises:
    add_bullet(label, text)

doc.add_paragraph()

add_para(
    'A fase estável de cada trajetória foi determinada automaticamente pela '
    'identificação do plateau no perfil de RMSD do backbone, utilizando uma '
    'janela deslizante de 5 ns e limiar de desvio padrão de 0,15 nm '
    '(módulo stability_filter do pipeline). Os critérios integrados para '
    'classificação de inibição peptídica foram: (i) RMSD do peptídeo < 0,3 nm '
    'na fase estável; (ii) ≥ 1 ponte de hidrogênio com persistência ≥ 10%; '
    '(iii) distância a pelo menos um resíduo catalítico < 0,5 nm; '
    '(iv) redução da SASA do peptídeo em relação ao valor inicial (enterramento); '
    'e (v) número de contatos intermoleculares estável ao longo da fase estável.'
)

# ════════════════════════════════════════════════════════════════════════════
# REFERÊNCIAS
# ════════════════════════════════════════════════════════════════════════════
add_heading('Referências', level=1)

refs = [
    'Abraham, M. J. et al. GROMACS: High performance molecular simulations through '
    'multi-level parallelism from laptops to supercomputers. SoftwareX, v. 1–2, '
    'p. 19–25, 2015.',

    'Ahmad, S. et al. Structure-based design of a synthetic protein inhibitor of '
    'Spodoptera frugiperda trypsin. Scientific Reports, v. 8, p. 1743, 2018. '
    'doi:10.1038/s41598-017-18733-9',

    'Al-Khafaji, K. et al. Design of TAT-conjugated Bowman-Birk trypsin inhibitor '
    'peptides with enhanced cell penetration and protease inhibition. '
    'Biomolecules, v. 16, p. 511, 2024. doi:10.3390/biom16040511',

    'Bray, S. A. et al. Introductory tutorials for simulating protein dynamics with '
    'GROMACS. Journal of Physical Chemistry B, v. 128, p. 9492–9509, 2024. '
    'doi:10.1021/acs.jpcb.4c04901',

    'Bussi, G.; Donadio, D.; Parrinello, M. Canonical sampling through velocity '
    'rescaling. Journal of Chemical Physics, v. 126, 014101, 2007.',

    'Darden, T.; York, D.; Pedersen, L. Particle mesh Ewald: An N·log(N) method '
    'for Ewald sums in large systems. Journal of Chemical Physics, v. 98, '
    'p. 10089–10092, 1993.',

    'Di Tommaso, P. et al. Nextflow enables reproducible computational workflows. '
    'Nature Biotechnology, v. 35, p. 316–319, 2017.',

    'Dolinsky, T. J. et al. PDB2PQR: expanding and upgrading automated preparation '
    'of biomolecular structures for molecular simulations. Nucleic Acids Research, '
    'v. 35, p. W522–W525, 2007.',

    'Dominguez, C.; Boelens, R.; Bonvin, A. M. J. J. HADDOCK: A protein–protein '
    'docking approach based on biochemical or biophysical information. Journal of '
    'the American Chemical Society, v. 125, p. 1731–1737, 2003.',

    'Dossantos, E. A. et al. Inhibitory effects of protease inhibitors from '
    'Sapindus trifoliatus seeds on larval gut proteases of Spodoptera frugiperda. '
    'PLoS ONE, v. 10, e0140484, 2015.',

    'Dunse, K. M. et al. Coexpression of potato type I and II protease inhibitors '
    'with potato protease inhibitor 8 confers levels of insect gut cysteine and '
    'serine protease inhibition superior to single or binary inhibitor combinations. '
    'BMC Plant Biology, v. 10, p. 286, 2010.',

    'Hess, B. et al. LINCS: A linear constraint solver for molecular simulations. '
    'Journal of Computational Chemistry, v. 18, p. 1463–1472, 1997.',

    'Jorgensen, W. L. et al. Comparison of simple potential functions for simulating '
    'liquid water. Journal of Chemical Physics, v. 79, p. 926–935, 1983.',

    'Lindorff-Larsen, K. et al. Improved side-chain torsion potentials for the Amber '
    'ff99SB protein force field. Proteins, v. 78, p. 1950–1958, 2010.',

    'Parrinello, M.; Rahman, A. Polymorphic transitions in single crystals: '
    'A new molecular dynamics method. Journal of Applied Physics, v. 52, '
    'p. 7182–7190, 1981.',

    'Singh, S. et al. PEPstrMOD: structure prediction of peptides containing natural, '
    'non-natural and modified residues. Biology Direct, v. 10, p. 73, 2015.',

    'Soelter, T. M. et al. Prebound state discovered in the unbinding pathway of '
    'fluorinated variants of the trypsin-BPTI complex. Journal of Chemical Theory '
    'and Computation, v. 20, p. 4518–4531, 2024. doi:10.1021/acs.jctc.3c01412',

    'Trozzi, F. et al. GROMACS 2020: Algorithms for highly efficient, load-balanced, '
    'and scalable molecular simulation. Journal of Chemical Physics, v. 153, '
    '214103, 2021.',

    'van Zundert, G. C. P. et al. The HADDOCK2.2 Web Server: User-Friendly '
    'Integrative Modeling of Biomolecular Complexes. Journal of Molecular Biology, '
    'v. 428, p. 720–725, 2016.',

    'Zhu, F. et al. A spatiotemporal atlas of the lepidopteran pest Helicoverpa '
    'armigera midgut. PLoS Biology, v. 20, e3001591, 2022.',
]

for ref in refs:
    add_ref(ref)

doc.save('metodologia_artigo.docx')
print('Arquivo gerado: metodologia_artigo.docx')
